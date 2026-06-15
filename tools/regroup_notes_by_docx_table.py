#!/usr/bin/env python3
from __future__ import annotations

import re
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


DOWNLOADS = Path("/Users/james/Downloads")
NOTES_ROOT = DOWNLOADS / "NOTES"
BACKUP_ROOT = NOTES_ROOT.with_name("NOTES.before_table_regroup")

DOC_TABLE_NAMES: dict[str, list[str]] = {
    "Big Data": [
        "Big Data",
    ],
    "Computer organisation": [
        "Internal hardware components of a computer",
        "The meaning of the stored program concept",
        "The processor and its components",
        "The Fetch-Execute cycle and the role of registers within it",
        "The processor instruction set",
        "Addressing modes",
        "Interrupts",
        "Factors affecting processor performance",
        "Input and output devices",
        "Secondary storage devices",
    ],
    "Computer systems": [
        "Hardware and software",
        "Classification of programming languages",
        "Types of program translator",
        "Logic gates",
        "Using Boolean algebra",
    ],
    "Data representation": [
        "Number systems",
        "Number bases",
        "Bits and bytes",
        "Units",
        "Unsigned binary",
        "Unsigned binary arithmetic",
        "Signed binary using two's complement",
        "Numbers with a fractional part",
        "Rounding errors",
        "Absolute and relative errors",
        "Normalisation of floating point form",
        "Underflow and overflow",
        "Information coding systems",
        "Analogue and digital",
        "Analogue digital conversion",
        "Bitmapped graphics",
        "Vector graphics",
        "Vector graphics versus bitmapped graphics",
        "Digital representation of sound",
        "MIDI",
        "Data compression",
        "Encryption",
    ],
    "Databases": [
        "Conceptual data models and entity relationship modelling",
        "Relational databases",
        "Database design and normalisation techniques",
        "Structured Query Language (SQL)",
        "Client server databases",
    ],
    "Functional programming": [
        "Function type",
        "First-class object",
        "Function application",
        "Partial function application",
        "Composition of functions",
        "Functional language programs",
        "Functional programming paradigm",
    ],
    "Networking": [
        "Communication methods",
        "Communication basics",
        "Network topology",
        "Types of networking between hosts",
        "Wireless networking",
        "The Internet and how it works",
        "Internet security",
        "TCP IP",
        "Standard application layer protocols",
        "IP address structure",
        "Subnet masking",
        "IP standards",
        "Public and private IP addresses",
        "Dynamic Host Configuration Protocol (DHCP)",
        "Network Address Translation (NAT)",
        "Port forwarding",
        "Client server model",
        "Thin- versus thick-client computing",
    ],
}

EXTRA_FOLDER_MERGES: dict[str, str] = {
    "Consequences of computing": "Individual moral social legal and cultural issues and opportunities",
    "Systematic approach": "Aspects of software development",
}

EXTRA_FOLDER_FOCUS: dict[str, str] = {
    "Consequences of computing": "Individual (moral), social (ethical), legal and cultural issues and opportunities.",
    "Systematic approach": "Analysis, design, implementation, testing and evaluation in software development.",
}

FOCUS_OVERRIDES: dict[str, list[str]] = {
    "Units": ["Know the names and symbols for units of information and powers of 2 and 10 used in file sizes."],
}


def clean_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = text.replace("\u00a0", " ")
    text = text.replace("→", "->")
    text = text.replace("à", "->")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    out: list[str] = []
    blank = False
    for line in lines:
        if not line:
            if out and not blank:
                out.append("")
            blank = True
            continue
        out.append(line)
        blank = False
    return "\n".join(out).strip()


def clean_note_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = text.replace("\u00a0", " ")
    text = text.replace("→", "->")
    text = text.replace("à", "->")
    out: list[str] = []
    blank = False
    for raw in text.splitlines():
        raw = raw.replace("\t", "    ").rstrip()
        indent_len = len(raw) - len(raw.lstrip(" "))
        indent = raw[:indent_len]
        body = re.sub(r"[ \t]+", " ", raw[indent_len:]).strip()
        if not body:
            if out and not blank:
                out.append("")
            blank = True
            continue
        out.append(indent + body)
        blank = False
    return "\n".join(out).strip()


def paragraph_indent(paragraph: Paragraph) -> tuple[int, bool]:
    ppr = paragraph._p.pPr
    if ppr is not None and ppr.numPr is not None and ppr.numPr.ilvl is not None:
        try:
            return min(8, max(0, int(ppr.numPr.ilvl.val) * 2)), True
        except Exception:
            return 0, True
    left = paragraph.paragraph_format.left_indent
    if left:
        return min(8, max(0, int(round(int(left) / 342900.0)) * 2)), False
    return 0, False


def paragraph_text(paragraph: Paragraph) -> str:
    text = clean_text(paragraph.text).replace("\n", " ")
    if not text:
        return ""
    spaces, is_list = paragraph_indent(paragraph)
    prefix = " " * spaces
    if is_list and not text.startswith(("- ", "* ")):
        return f"{prefix}- {text}"
    return prefix + text


def slug_safe(name: str) -> str:
    name = re.sub(r'[<>:"/\\|?*]+', " ", name)
    name = re.sub(r"\s+", " ", name).strip()
    return name or "Notes"


GENERIC_HEADINGS = {
    "BELOW",
    "COMPARISON",
    "DEFINITION",
    "EXAMPLE",
    "EXAMPLES",
    "PPQ",
    "PURPOSE",
    "RECAP",
    "CHARACTERISTICS",
}

LABEL_HEADINGS = {
    "ADVANTAGES",
    "ADVANTAGES AND DISADVANTAGES",
    "CHARACTERISTICS",
    "COMPARISON",
    "CONVERTING TO TWO'S COMPLEMENT",
    "DATA FROM",
    "DEFINITION",
    "DEFINITIONS",
    "DISADVANTAGES",
    "EXAMPLE",
    "EXAMPLES",
    "HOW DOES IT WORK",
    "HOW DOES IT WORK?",
    "HOW TO INCREASE PRECISION?",
    "MAIN FEATURES",
    "NOTE",
    "PRINCIPLES OF OPERATION",
    "PROS & CONS",
    "PURPOSE",
    "PURPOSE?",
    "PURPOSES",
    "RECAP",
    "RANGE",
    "SENDING DATA",
    "SUITABILITY",
    "TYPES",
    "USES",
    "WHERE USED?",
    "WHY?",
}


def iter_block_items(parent):
    element = getattr(parent, "element", None)
    container = element.body if element is not None and hasattr(element, "body") else parent._tc
    for child in container.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def markdown_table(table: Table) -> list[str]:
    rows: list[list[str]] = []
    max_cols = 0
    for row in table.rows:
        cells = [clean_text(cell_text(cell, include_nested=False)).replace("\n", "; ") for cell in row.cells]
        max_cols = max(max_cols, len(cells))
        rows.append(cells)
    if not rows or max_cols == 0:
        return []
    rows = [r + [""] * (max_cols - len(r)) for r in rows]
    def esc(s: str) -> str:
        return s.replace("|", "/").strip()
    out = ["| " + " | ".join(esc(c) for c in rows[0]) + " |"]
    out.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    for row in rows[1:]:
        out.append("| " + " | ".join(esc(c) for c in row) + " |")
    return out


def cell_text(cell: _Cell, include_nested: bool = True) -> str:
    parts: list[str] = []
    for block in iter_block_items(cell):
        if isinstance(block, Paragraph):
            text = paragraph_text(block)
            if text:
                parts.append(text)
        elif include_nested:
            table_lines = markdown_table(block)
            if table_lines:
                parts.extend(table_lines)
    return "\n".join(parts)


def short_heading(text: str, fallback: str) -> str:
    text = clean_text(text)
    if not text:
        return fallback
    first = text.splitlines()[0]
    first = re.sub(r"^(Know and understand how to|Know that|Know how to|Understand|Be able to|Be familiar with|Explain|Describe|Compare and contrast)\b", "", first, flags=re.I)
    first = re.sub(r"[^A-Za-z0-9 +/&-]+", " ", first)
    stop = {
        "a", "an", "and", "are", "as", "be", "between", "by", "for", "from", "in", "is",
        "it", "of", "on", "some", "that", "the", "to", "with", "following", "about",
    }
    words = [w for w in first.split() if w.lower() not in stop]
    if not words:
        return fallback
    return " ".join(words[:7]).title()


def heading_from_cell(text: str, fallback: str) -> str:
    for line in clean_text(text).splitlines():
        stripped = line.strip(" :-")
        if len(stripped) < 3:
            continue
        if stripped.upper() in GENERIC_HEADINGS or re.fullmatch(r"EXAMPLE\s*\d*", stripped.upper()):
            continue
        if stripped.upper() == stripped and re.search(r"[A-Z]", stripped):
            return stripped.title()
        break
    return fallback


def heading_from_row(row, detail_col: int, fallback: str) -> str:
    for i in range(detail_col - 1, -1, -1):
        text = clean_text(row.cells[i].text)
        if text and not re.fullmatch(r"\d+(\.\d+)*", text):
            return short_heading(text, fallback)
    return fallback


def is_heading_label(line: str) -> bool:
    stripped = line.strip().strip(":")
    upper = stripped.upper()
    if upper in LABEL_HEADINGS:
        return True
    if line.endswith(":") and len(stripped) <= 70 and len(stripped.split()) <= 7 and not re.search(r"[=<>]", stripped):
        return True
    if upper == stripped and 3 <= len(stripped) <= 70 and re.search(r"[A-Z]", stripped) and not re.search(r"[=<>]", stripped):
        return True
    return False


def title_label(line: str) -> str:
    return re.sub(r"\s+", " ", line.strip().strip(":")).title()


def is_code_or_formula(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    upper = s.upper().strip(":")
    if upper in LABEL_HEADINGS or (s.endswith(":") and len(s.split()) <= 7):
        return False
    if s in {"+", "-", "*", "/", "(", ")", "&&", "||"}:
        return True
    if re.fullmatch(r"-?\d+(\.\d+)?", s) or s in {"True", "False"}:
        return True
    if s.startswith(("ghci>", "```")):
        return True
    if s in {"DELETE", ")"}:
        return True
    if re.match(r"^SELECT\s+(<|\*)", s, re.I):
        return True
    if re.match(r"^FROM\s+<", s, re.I):
        return True
    if re.match(r"^WHERE\s+(<|\()", s, re.I):
        return True
    if re.match(r"^ORDER BY\s+<", s, re.I):
        return True
    if re.match(r"^(CREATE TABLE|PRIMARY KEY|INSERT INTO|VALUES|UPDATE|SET)\s*(<|\()", s, re.I):
        return True
    if re.match(r"^[a-z][A-Za-z0-9_]*(\s+[A-Za-z_][A-Za-z0-9_]*)*\s*=", s):
        return True
    if re.search(r"[<>{}]|::|->|\bmod\b", s):
        return True
    if re.match(r"^[01](\s*[+*]\s*[01])+", s):
        return True
    if re.match(r"^-\d", s):
        return True
    return False


def should_bullet(line: str) -> bool:
    s = line.strip()
    if s in {"Create", "Retrieve", "Update", "Delete"}:
        return True
    if not s or s.startswith(("-", "#")) or s.startswith("| ") or is_code_or_formula(s):
        return False
    return True


def format_note_body(text: str) -> str:
    out: list[str] = []
    for line in clean_note_text(text).splitlines():
        indent_len = len(line) - len(line.lstrip(" "))
        indent = line[:indent_len]
        s = line.strip()
        if not s:
            if out and out[-1] != "":
                out.append("")
            continue
        if s.startswith("| ") and s.count("|") >= 2:
            if out and out[-1] != "" and not out[-1].startswith("|"):
                out.append("")
            out.append(s)
            continue
        if s.startswith(("- ", "* ")):
            out.append(f"{indent}- {s[2:].strip()}")
            continue
        if is_code_or_formula(s) or indent_len >= 4:
            out.append(indent + s)
            continue
        if is_heading_label(s):
            if out and out[-1] != "":
                out.append("")
            out.append(f"### {title_label(s)}")
            out.append("")
            continue
        if should_bullet(s):
            out.append(f"{indent}- {s}")
        else:
            out.append(indent + s)
    return clean_note_text("\n".join(out))


def remove_empty_level3_headings(text: str) -> str:
    lines = text.splitlines()
    keep = [True] * len(lines)
    for i, line in enumerate(lines):
        if not line.startswith("### "):
            continue
        j = i + 1
        while j < len(lines) and not lines[j].strip():
            j += 1
        if j >= len(lines) or lines[j].startswith("#"):
            keep[i] = False
            k = i + 1
            while k < len(lines) and not lines[k].strip():
                keep[k] = False
                k += 1
    return clean_note_text("\n".join(line for line, ok in zip(lines, keep) if ok))


def table_detail_column(table) -> int:
    if not table.rows:
        return 0
    header = [clean_text(cell.text).lower() for cell in table.rows[0].cells]
    for i in range(len(header) - 1, -1, -1):
        h = header[i]
        if h and "syllabus" not in h and "syll ref" not in h:
            return i
    return len(header) - 1


def table_syllabus_column(table) -> int:
    if not table.rows:
        return -1
    header = [clean_text(cell.text).lower() for cell in table.rows[0].cells]
    for i, h in enumerate(header):
        if "syllabus content" in h:
            return i
    return 0 if len(header) > 1 else -1


def is_repeated_header_row(row) -> bool:
    cells = [clean_text(cell.text).lower() for cell in row.cells]
    return bool(cells) and "syllabus content" in cells[0] and any("notes" in c for c in cells[1:])


def extract_doc_tables(doc_path: Path, table_names: list[str]) -> list[dict[str, str]]:
    doc = Document(doc_path)
    if len(doc.tables) != len(table_names):
        raise RuntimeError(f"{doc_path.name}: expected {len(table_names)} tables, found {len(doc.tables)}")
    result: list[dict[str, str]] = []
    for idx, (table, table_name) in enumerate(zip(doc.tables, table_names), 1):
        detail_col = table_detail_column(table)
        sections: list[str] = [f"# {table_name}"]
        data_rows = [
            row for row in table.rows[1:]
            if not is_repeated_header_row(row)
            and detail_col < len(row.cells)
            and clean_text(cell_text(row.cells[detail_col]))
        ]
        include_row_headings = len(data_rows) > 1
        for row in data_rows:
            if detail_col >= len(row.cells):
                continue
            text = format_note_body(cell_text(row.cells[detail_col]))
            if not text:
                continue
            row_fallback = ""
            title = heading_from_cell(text, row_fallback)
            if include_row_headings and title and not text.lstrip().startswith("###"):
                sections.append("")
                sections.append(f"## {title}")
                sections.append("")
            elif sections[-1] != "":
                sections.append("")
            sections.append(text)
        result.append({"name": table_name, "text": remove_empty_level3_headings(clean_note_text("\n".join(sections))) + "\n"})
    return result


def tokens(text: str) -> set[str]:
    return {t for t in re.findall(r"[a-z0-9]+", text.lower()) if len(t) > 2}


def score_table(old_name: str, old_text: str, table: dict[str, str]) -> int:
    stem_tokens = tokens(old_name)
    table_tokens = tokens(table["name"] + "\n" + table["text"])
    score = len(stem_tokens & table_tokens) * 4
    norm_stem = re.sub(r"[^a-z0-9]+", "", old_name.lower())
    norm_table = re.sub(r"[^a-z0-9]+", "", table["text"].lower())
    if norm_stem and norm_stem in norm_table:
        score += 30
    first_heading = ""
    for line in old_text.splitlines():
        if line.startswith("# "):
            first_heading = line[2:].strip()
            break
    if first_heading and re.sub(r"[^a-z0-9]+", "", first_heading.lower()) in norm_table:
        score += 30
    return score


def exam_notes_section(text: str) -> str:
    marker = "## EXAM NOTES"
    idx = text.find(marker)
    if idx >= 0:
        rest = text[idx + len(marker):]
        next_heading = re.search(r"\n##\s+", rest)
        if next_heading:
            rest = rest[:next_heading.start()]
        return clean_text(rest)
    return ""


def supplementary_section(path: Path, text: str, table_text: str) -> str:
    notes = exam_notes_section(text)
    if not notes and path.stem == "Haskell Paper Notes":
        notes = clean_text(text)
    if not notes:
        return ""
    norm_notes = re.sub(r"\s+", " ", notes.lower())
    norm_table = re.sub(r"\s+", " ", table_text.lower())
    if norm_notes in norm_table:
        return ""
    return notes


def normalize_heading(text: str) -> str:
    text = re.sub(r"^\s*the\s+", "", text.lower())
    return re.sub(r"[^a-z0-9]+", "", text)


def insert_under_heading(text: str, heading: str, body: str) -> str:
    if not body:
        return text
    lines = text.rstrip().splitlines()
    target = normalize_heading(heading)
    best = None
    for i, line in enumerate(lines):
        if not line.startswith("#"):
            continue
        title = line.lstrip("#").strip()
        if normalize_heading(title) == target:
            best = i
            break
    if best is None:
        return clean_text(text.rstrip() + f"\n\n### {heading}\n\n{body}") + "\n"

    level = len(lines[best]) - len(lines[best].lstrip("#"))
    end = len(lines)
    for j in range(best + 1, len(lines)):
        if not lines[j].startswith("#"):
            continue
        other_level = len(lines[j]) - len(lines[j].lstrip("#"))
        if other_level <= level:
            end = j
            break
    existing = "\n".join(lines[best:end])
    if re.sub(r"\s+", " ", body.lower()) in re.sub(r"\s+", " ", existing.lower()):
        return text
    insert = [""] + body.splitlines() + [""]
    lines[end:end] = insert
    return clean_text("\n".join(lines)) + "\n"


def merge_supplements(text: str, supplements: list[tuple[str, str]]) -> str:
    for heading, body in supplements:
        text = insert_under_heading(text, heading, body)
    return text


def strip_syllabus_focus_sections(text: str) -> str:
    lines: list[str] = []
    skipping = False
    for line in text.splitlines():
        stripped = line.strip()
        if stripped == "## Syllabus focus":
            skipping = True
            continue
        if skipping and stripped.startswith("#"):
            skipping = False
        if not skipping:
            lines.append(line)
    return clean_note_text("\n".join(lines)) + "\n"


def strip_exam_note_headings(text: str) -> str:
    lines = []
    for line in text.splitlines():
        if line.strip() in {"## EXAM NOTES", "## EXTRA EXAM NOTES"}:
            continue
        lines.append(line)
    return strip_syllabus_focus_sections("\n".join(lines))


def replace_once(text: str, old: str, new: str) -> str:
    if old not in text:
        raise RuntimeError(f"manual note update target not found: {old[:60]!r}")
    return text.replace(old, new, 1)


def replace_section(text: str, start_heading: str, end_heading: str, new: str) -> str:
    start = text.find(start_heading)
    if start < 0:
        raise RuntimeError(f"manual note section start not found: {start_heading!r}")
    end = text.find(end_heading, start + len(start_heading))
    if end < 0:
        raise RuntimeError(f"manual note section end not found: {end_heading!r}")
    return text[:start] + new.rstrip() + "\n\n" + text[end:]


def apply_manual_updates(doc_name: str, table_name: str, text: str) -> str:
    if doc_name == "Computer organisation" and table_name == "Secondary storage devices":
        old = """### Principles Of Operation

- A laser beam is focused onto the surface of the disk
- Data is stored in the form of microscopic pits and lands
- The pits and lands represent binary data (0s and 1s)
- The laser beam detects these variations to read the information stored on the disk
- When reading data, the laser beam reflects off the surface of the disk, and the reflected light is detected by a sensor
- When writing data, the laser beam alters the surface of the disk to create the pits and lands
- The optical drive contains:
  - a motor to spin the disk
  - a lens system to focus the laser beam accurately on the data tracks"""
        new = """### Optical Disk Reading Mechanism

- A low-power laser/light is shone at the disk.
- The light is focused onto a spot on the track.
- The disk spins as the track is read.
- The disk may use constant linear velocity, zoned constant linear velocity, or variable angular velocity.
- Some light is reflected back from the disk.
- A light sensor detects/measures the amount of reflected light.
- The reflected-light pattern is converted back into binary data.

### How Optical Data Is Represented

- Data is stored on one spiral track.
- Data is stored using pits and lands/marks with different reflective properties.
- A continuation of a pit or land reflects light in the same way.
- A transition between pit and land scatters or changes the reflected light.
- A transition between pit and land represents 1.
- A continuation of a pit or land represents 0.

### Optical Disk Writing Mechanism

- A higher-power laser is used when writing/burning data.
- On recordable disks, the laser heats/burns/changes a dye or recording layer.
- This changes the reflectivity of the surface to create marks that can be read like pits and lands.
- On rewritable disks, the laser changes the state/phase of the material.
- The different reflective states are used to represent binary data.

### The Optical Drive Contains

- a motor to spin the disk
- a lens system to focus the laser beam accurately on the data tracks"""
        text = replace_once(text, old, new)
        text = text.replace("\nno moving parts = more durable + less prone to mechanical failure", "\n- no moving parts = more durable + less prone to mechanical failure")
        return text

    if doc_name == "Networking" and table_name == "Internet security":
        new = """### Asymmetric Encryption

- Uses a mathematically linked public/private key pair.
- Each communicating device has its own public key and private key.
- Public keys can be shared openly.
- Private keys must be kept secret by their owner.
- Data encrypted with a public key can only be decrypted with the matching private key.
- Data encrypted with a private key can be checked/decrypted with the matching public key.
- This helps solve the symmetric-key exchange problem, but it is slower than symmetric encryption.

### Asymmetric Encryption For Confidentiality

- The receiver generates a public/private key pair.
- The receiver shares their public key, often using a digital certificate.
- The sender checks the certificate to confirm the public key belongs to the receiver.
- The sender encrypts the plaintext, or a symmetric session key, using the receiver's public key.
- The ciphertext is sent to the receiver.
- The receiver decrypts it using the receiver's private key.
- Only the receiver should be able to decrypt it because only the receiver has that private key.

### Digital Signature Step-By-Step

- The sender applies a hashing algorithm to the plaintext to create a message digest.
- The sender encrypts the digest with the sender's private key.
- The encrypted digest is the digital signature.
- The sender appends the digital signature to the plaintext.
- If confidentiality is also needed, the sender encrypts the plaintext and signature using the receiver's public key.
- The receiver decrypts the ciphertext using the receiver's private key.
- The receiver now has the plaintext and the sender's encrypted digest.
- The receiver decrypts the digital signature using the sender's public key.
- This gives the digest that the sender created.
- The receiver hashes the received plaintext using the same hashing algorithm.
- The receiver compares the two digests.
- If the digests match, the message has not been changed and the sender is authenticated.
- If the digests do not match, the message may have been altered/corrupted or the signature may not be from the sender.

### Digital Certificates

- A digital certificate verifies ownership of a public key.
- It helps prove that a public key belongs to the claimed person/organisation.
- It helps prevent an attacker using a fake public key.
- Certificates are issued by certificate authorities after checking identity/key ownership.
- A browser/device checks the certificate authority's signature to decide whether to trust the certificate.

### Contains

- serial number
- owner's name/identity
- expiry date
- owner's public key
- certificate authority's digital signature"""
        return replace_section(text, "### Asymmetric Encryption", "### Worms", new)

    if doc_name == "Data representation" and table_name == "Encryption":
        text = text.replace("ciphertext is rando, so", "ciphertext is random, so")
        insert = """- Asymmetric encryption and digital signatures are covered in more detail in Networking > Internet security."""
        if insert not in text:
            text = text.rstrip() + "\n" + insert + "\n"
    return text


def polish_note_text(text: str) -> str:
    replacements = {
        " cant ": " cannot ",
        " cant\n": " cannot\n",
        "cant conform": "cannot conform",
        "cant change": "cannot change",
        "facts cant": "facts cannot",
        "you cant": "you cannot",
        "it cant": "it cannot",
        "that cant": "that cannot",
        "they cant": "they cannot",
        " doesnt ": " does not ",
        " doesnt\n": " does not\n",
        " isnt ": " is not ",
        " isnt\n": " is not\n",
        " dont ": " do not ",
        " dont\n": " do not\n",
        " ie ": " i.e. ",
        "(ie ": "(i.e. ",
        " eg ": " e.g. ",
        ", eg ": ", e.g. ",
        " num of ": " number of ",
        " ppl ": " people ",
        " cipher text": " ciphertext",
        "Cipher text": "Ciphertext",
        "malwares": "malware",
        "companies servers": "company's servers",
        "vulnerably/improvements fixed": "vulnerabilities or improvements fixed",
        "dark drives": "USB/removable drives",
        "story up to": "store up to",
        "time requires": "time required",
        "accepted the offer": "accepts the offer",
        "internationaly": "internationally",
        "ciphertext is rando,": "ciphertext is random,",
        "the have no side effects": "they have no side effects",
        "takes two integers as arguments, nd returns": "takes two integers as arguments, and returns",
        "Parenthesis are required": "Parentheses are required",
        "Pre-pending": "Prepending",
        "same as python": "same as Python",
        "If you can touch it, its hardware": "If you can touch it, it is hardware",
        "NOTE: its exactly": "NOTE: it is exactly",
        "NOTE: its not enough": "NOTE: it is not enough",
        "If they intercept the key when its shared": "If they intercept the key when it is shared",
        "Sound might not be faithfully represented if its less than this": "Sound might not be faithfully represented if the sampling frequency is less than twice the maximum signal frequency",
        "Interpreter translates a program each time its executed": "Interpreter translates a program each time it is executed",
        "when a file is opened/downloaded, it checks it against its known database of malware": "when a file is opened/downloaded, anti-malware checks it against its known database of malware",
        "malware files are quarantined to prevent it from spreading": "malware files are quarantined to prevent them from spreading",
        "on an ongoing bases": "on an ongoing basis",
        "Its a design methodology": "It is a design methodology",
        "its not enough": "it is not enough",
        "its not very durable": "it is not very durable",
        "its assumed": "it is assumed",
        "its less": "it is less",
        "applies repeatedly (using recursion) that combining function": "repeatedly applies that combining function",
        "Then....": "Then:",
        "....Practise, practise, practise": "Practise map, filter and fold examples until the trace can be written on paper.",
        "...which means...": "This means:",
        "...taking one argument": "It takes one argument",
        "...and ++": "Use ++",
        "...but contents": "But the contents",
        "The set of real numbers // the set of all possible real-world quantities;": "The set of real numbers / the set of all possible real-world quantities.",
        "<br>": "; ",
        "<BR>": "; ",
        "\n- - *\n": "\n- *\n",
        "aand": "and",
        "randomm": "random",
        "copied of the shared key": "copy of the shared key",
        "Wont": "Will Not",
        "wont": "will not",
        "cypher": "cipher",
        "number of its": "number of bits",
        "specified num ": "specified number ",
        "specified num\n": "specified number\n",
        "numberber": "number",
        "Its an acronym": "It is an acronym",
        "Its theoretically": "It is theoretically",
        "In routing, its changed": "In routing, it is changed",
        "shares the name network identifier": "shares the same network identifier",
        "BUT vulnerable to interception …": "But the key exchange is vulnerable to interception.",
        "SOLUTIONS TO THE MALWARE's LISTED ABOVE": "Ways to reduce malware risk:",
        "in regards to": "with regard to",
        "Higher order function = A function": "Higher-order function = a function",
        " // returns a function as a result // ": " / returns a function as a result / ",
        "This is why add function type is integer --> integer --> integer": "- This is why add function type is integer --> integer --> integer",
        "If f: A -> B, we say that the function type is A -> B where A is the argument type and B is the result type": "- If f: A -> B, we say that the function type is A -> B where A is the argument type and B is the result type",
        "domain = set that a function takes input values from": "- domain = set that a function takes input values from",
        "Invert all bits, making 0 -> 1 and 1 -> 0": "- Invert all bits, making 0 -> 1 and 1 -> 0",
        "The type of this function is f: integer x integer --> integer, where integer x integer is the cartesian product of set Integer with itself": "- The type of this function is f: integer x integer --> integer, where integer x integer is the cartesian product of set Integer with itself",
        "Describe the co-domain of function f: N --> R (PPQ)": "- PPQ: describe the co-domain of function f: N --> R",
        "Eg if f : A --> B, and g : B --> C, then": "- E.g. if f : A --> B, and g : B --> C, then",
        "It takes one argument after another, and returning": "It takes one argument after another and returns",
        "### Eg": "### Example",
        "- Eg ": "- E.g. ",
        "Eg ": "E.g. ",
        "CONVERTING FLOATING POINT BINARY -> DECIMAL": "\n### Converting Floating Point Binary To Decimal\n",
        "same as Python…": "same as Python:",
        "It prints …": "It prints:",
        "ALTHOUGH there are different methods of representation. See the examples below…": "There are different methods of representation. See the examples below.",
        "sounds like a lot, but it is not enough… hence IPv6": "sounds like a lot, but it is not enough; hence IPv6",
        "just think about it….": "Key idea: compare the size of the error with the size of the original value.",
        "Minimises data duplication (don’t say eliminates…)": "Minimises data duplication; do not say it eliminates all duplication.",
        "On mac, select an item in finder and press cmd+i …": "On macOS, select an item in Finder and press Cmd+I.",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)

    line_replacements = {
        "- A. During execute phase MBR used to store other data": "- Accepted wording: during the execute phase, the MBR may be used to store other data.",
        "- A. Further instructions may need to be fetched before the instruction has finished executing, if pipelining/parallelisation is referenced explicitly in the response": "- Accepted wording: further instructions may be fetched before an instruction has finished executing if pipelining/parallelisation is explicitly referenced.",
        "- A. MBR is not (directly) wired to the (processor) components that will execute the instruction which CIR is": "- Accepted wording: the MBR is not directly wired to the processor components that execute the instruction, but the CIR is.",
        "- A. The MBR is not (directly) wired to the ALU as BOD": "- Accepted wording: the MBR is not directly wired to the ALU, unlike the CIR/processor execution path.",
        "- A. programs can be moved in to (and out of) main memory": "- Accepted wording: programs can be moved into and out of main memory.",
        "- NE. So programs/tasks will run faster": "- Not enough: simply saying programs/tasks will run faster.",
        "- NE. More efficient": "- Not enough: simply saying it is more efficient.",
        "- A. real numbers": "- Accepted wording: real numbers.",
        "- A. numbers that represent any quantity along an infinite number line": "- Accepted wording: numbers that represent quantities on an infinite number line.",
        "- A. all numbers excluding imaginary/complex numbers": "- Accepted wording: all numbers excluding imaginary/complex numbers.",
        "- A. rational and irrational numbers": "- Accepted wording: rational and irrational numbers.",
        "- A. “Parameter”, “Input” for “Argument”": "- Accepted wording: \"parameter\" or \"input\" can be used for \"argument\".",
        "- NE. A function that uses another function": "- Not enough: a function that only uses/calls another function.",
        "- R. Explanations that are specifically of the map function": "- Reject: explanations that are only about the map function.",
        "- A. DHCP server / router would need to be configured to allocate a fixed / static IP address to the web server": "- Accepted wording: configure the DHCP server/router to allocate a fixed/static IP address to the web server.",
        "- A. just \"IP\" for \"IP address\" as BOD": "- Accepted wording: \"IP\" can be used for \"IP address\" if the meaning is clear.",
        "- NE. there are more IPv6 addresses": "- Not enough: simply saying there are more IPv6 addresses.",
        "Higher word length -> more bits can be transferred as a single unit": "- Higher word length -> more bits can be transferred as a single unit",
        "Data can be accessed more quickly -> programs run faster": "- Data can be accessed more quickly -> programs run faster",
        "Higher clock speed -> more FE cycles carried out per second -> faster processing of tasks": "- Higher clock speed -> more FE cycles carried out per second -> faster processing of tasks",
        "Higher clock speed -> higher power consumption (overheating? -> reduce performance)": "- Higher clock speed -> higher power consumption, which can cause overheating and reduce performance",
        "More unique memory locations -> reduces need for secondary storage / virtual memory (these are slower)": "- More unique memory locations -> reduces need for secondary storage / virtual memory, which are slower",
        "Wider data bus -> less FE cycles needed to fetch instructions from memory -> faster executions": "- Wider data bus -> fewer FE cycles needed to fetch instructions from memory -> faster execution",
        "Exam questions can ask about current individual, social, legal and cultural opportunities and risks of computing.": "- Exam questions can ask about current individual, social, legal and cultural opportunities and risks of computing.",
        "Co-domain = the outputs from a function <-- known as the range in maths": "- Co-domain = the outputs from a function; known as the range in maths",
        "NOT portable -> they are processor specific": "- NOT portable -> they are processor specific",
        "CRUD -> what you want to do to data": "- CRUD -> what you want to do to data",
        "REST -> how you do it over a network": "- REST -> how you do it over a network",
        "Public IP -> private IP": "- Public IP -> private IP",
        "Public IP -> private IP of server": "- Public IP -> private IP of server",
        "To correctly draw the ERD from a given standard form database -> FOLLOW THE PRIMARY KEY; this means, starting from a table, where else does its primary key appear? When found - that's the FK, and that’s a one-to-many relationship": "- To correctly draw the ERD from a given standard form database -> follow the primary key; starting from a table, find where else its primary key appears. When found, that is the foreign key and shows a one-to-many relationship.",
        "## Big Data Catch-All Term Data Will Not Fit": "## Big Data That Does Not Fit Traditional Processing",
        "## CORE IDEA": "## Core Idea",
        "## ANALYSIS": "## Analysis",
        "## DESIGN": "## Design",
        "## IMPLEMENTATION": "## Implementation",
        "## TESTING": "## Testing",
        "## EVALUATION": "## Evaluation",
        "## Compare Capacity Speed Access Various Media Make": "## Comparing Secondary Storage Media",
        "## Differentiate Character Code Representation Decimal Digit Its": "## Character Code Digits Versus Binary Values",
        "## Compare Absolute Relative Errors Large Small Magnitude": "## Comparing Absolute And Relative Errors",
        "## 2 N Different Values Can Represented N": "## Number Of Values Represented By N Bits",
        "## How Bitmaps Represented": "## How Bitmaps Are Represented",
        "## Calculate Storage Requirements Bitmapped Images Aware Bitmap": "## Calculating Bitmap Storage Requirements",
        "## How Vector Graphics Represent Images Using Lists": "## How Vector Graphics Represent Images",
        "## Use Vector Graphic Primitives Create Simple Vector": "## Using Vector Graphic Primitives",
        "## Digital Representation Sound Terms": "## Digital Sound Terms",
        "## Calculate Sound Sample Sizes Bytes": "## Calculating Sound Sample Sizes",
        "## What Meant Encryption Able Define": "## Encryption Definition",
        "## Compare Vernam Cipher Ciphers Depend Computational Security": "## Vernam Cipher Versus Computational Security",
        "## Term Imperative High-Level Language Its Relationship Low-Level": "## Imperative High-Level Languages And Low-Level Languages",
        "## Role Four Layers Tcp/Ip Stack Application Transport": "## TCP/IP Stack Layers",
        "## Discuss How Improved Code Quality Monitoring Protection": "## Improved Code Quality Monitoring And Malware Protection",
    }
    lines = []
    for line in text.splitlines():
        stripped = line.strip()
        indent = line[:len(line) - len(line.lstrip(" "))]
        if stripped == "- - *":
            lines.append(indent + "- *")
            continue
        replacement = line_replacements.get(line)
        if replacement is None and stripped in line_replacements:
            replacement = line_replacements[stripped]
            if indent and replacement.startswith("- "):
                replacement = indent + replacement
        lines.append(replacement if replacement is not None else line)
    text = "\n".join(lines)

    text = text.replace(
        "### Ne. ℝ\n\n- TO. real number stated but then another set described",
        "- Do not only write: R/real numbers without explaining the set.\n- Avoid mixing real numbers with another different set in the same answer.",
    )
    text = text.replace(
        "- ensures any code vulnerabilities or improvements fixed by the developer are applied to the app",
        "- ensures security fixes and improvements made by the developer are applied to the app",
    )
    text = re.sub(r"(?m)^- - \*$", "- -\n- *", text)
    text = text.replace("numberber", "number")
    text = re.sub(r"(?m)^([ \t]*-[ \t]+)-[ \t]+", r"\1", text)
    return strip_syllabus_focus_sections(text)


def drop_first_h1(text: str) -> str:
    lines = text.splitlines()
    if lines and lines[0].startswith("# "):
        lines = lines[1:]
        while lines and not lines[0].strip():
            lines = lines[1:]
    return clean_text("\n".join(lines))


def merge_extra_folder(folder: Path, title: str) -> str:
    sections = [f"# {title}"]
    for txt in sorted(folder.glob("*.txt")):
        body = strip_exam_note_headings(txt.read_text(encoding="utf-8", errors="ignore"))
        body = drop_first_h1(body)
        if not body:
            continue
        sections.extend(["", f"## {txt.stem}", "", body])
    return clean_text("\n".join(sections)) + "\n"


def build_new_notes() -> tuple[Path, dict[str, int]]:
    tmp = Path(tempfile.mkdtemp(prefix="notes_table_regroup_"))
    counts: dict[str, int] = {}
    for doc_name, table_names in DOC_TABLE_NAMES.items():
        doc_path = DOWNLOADS / f"{doc_name}.docx"
        if not doc_path.exists():
            raise FileNotFoundError(doc_path)
        out_dir = tmp / doc_name
        out_dir.mkdir(parents=True, exist_ok=True)
        tables = extract_doc_tables(doc_path, table_names)
        source_root = BACKUP_ROOT if BACKUP_ROOT.exists() else NOTES_ROOT
        old_files = sorted((source_root / doc_name).glob("*.txt")) if (source_root / doc_name).exists() else []
        assigned: dict[int, list[tuple[str, str]]] = {i: [] for i in range(len(tables))}
        for old in old_files:
            old_text = old.read_text(encoding="utf-8", errors="ignore")
            best = max(range(len(tables)), key=lambda i: score_table(old.stem, old_text, tables[i]))
            extra = supplementary_section(old, old_text, tables[best]["text"])
            if extra:
                assigned[best].append((old.stem, extra))
        for i, table in enumerate(tables):
            text = table["text"].rstrip()
            if assigned[i]:
                seen = set()
                unique = []
                for heading, body in assigned[i]:
                    key = (heading, body)
                    if key not in seen:
                        unique.append(key)
                        seen.add(key)
                text = merge_supplements(text, unique).rstrip()
            text = polish_note_text(apply_manual_updates(doc_name, table["name"], text)).rstrip()
            (out_dir / f"{slug_safe(table['name'])}.txt").write_text(text.rstrip() + "\n", encoding="utf-8")
        counts[doc_name] = len(tables)

    extra_source_root = BACKUP_ROOT if BACKUP_ROOT.exists() else NOTES_ROOT
    for folder in extra_source_root.iterdir():
        if folder.is_dir() and folder.name not in DOC_TABLE_NAMES:
            if folder.name in EXTRA_FOLDER_MERGES:
                out_folder = tmp / folder.name
                out_folder.mkdir(parents=True, exist_ok=True)
                title = EXTRA_FOLDER_MERGES[folder.name]
                (out_folder / f"{slug_safe(title)}.txt").write_text(polish_note_text(merge_extra_folder(folder, title)), encoding="utf-8")
                continue
            out_folder = tmp / folder.name
            shutil.copytree(folder, out_folder, dirs_exist_ok=True)
            for txt in out_folder.rglob("*.txt"):
                txt.write_text(polish_note_text(strip_exam_note_headings(txt.read_text(encoding="utf-8", errors="ignore"))), encoding="utf-8")
    return tmp, counts


def audit(root: Path) -> None:
    files = list(root.rglob("*.txt"))
    if not files:
        raise RuntimeError("no generated txt files")
    for p in files:
        data = p.read_text(encoding="utf-8", errors="ignore")
        if "\r" in data:
            raise RuntimeError(f"CR line ending: {p}")
        if "\n\n\n" in data:
            raise RuntimeError(f"triple blank line: {p}")
        for i, line in enumerate(data.splitlines(), 1):
            if line.rstrip() != line:
                raise RuntimeError(f"trailing whitespace: {p}:{i}")


def main() -> None:
    new_root, counts = build_new_notes()
    audit(new_root)
    if NOTES_ROOT.exists():
        if not BACKUP_ROOT.exists():
            shutil.copytree(NOTES_ROOT, BACKUP_ROOT)
        shutil.rmtree(NOTES_ROOT)
    shutil.copytree(new_root, NOTES_ROOT)
    shutil.rmtree(new_root)
    audit(NOTES_ROOT)
    total = len(list(NOTES_ROOT.rglob("*.txt")))
    print(f"wrote {total} txt files")
    for name in sorted(counts):
        print(f"{name}: {counts[name]} table files")
    print(f"backup: {BACKUP_ROOT}")


if __name__ == "__main__":
    main()
