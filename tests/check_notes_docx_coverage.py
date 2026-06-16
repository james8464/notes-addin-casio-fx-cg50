#!/usr/bin/env python3
from __future__ import annotations

import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ModuleNotFoundError:
    raise SystemExit("FAIL notes docx coverage: python-docx unavailable")

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
from tools.regroup_notes_by_docx_table import (  # noqa: E402
    DOC_TABLE_NAMES,
    clean_text,
    iter_block_items,
    markdown_table,
    slug_safe,
    table_detail_column,
)

NOTES = ROOT / "calculator_files" / "NOTES"
DOCX_DIR = Path(os.environ.get("NOTES_DOCX_DIR", "/Users/james/Downloads"))
DOCX_NAMES = [
    "Big Data.docx",
    "Computer organisation.docx",
    "Computer systems.docx",
    "Data representation.docx",
    "Databases.docx",
    "Functional programming.docx",
    "Networking.docx",
]


def normalise(text: str) -> str:
    text = (
        text.replace("\u00a0", " ")
        .replace("e.g.", "eg")
        .replace("E.g.", "eg")
        .replace("i.e.", "ie")
        .replace("I.e.", "ie")
        .replace("can't", "cant")
        .replace("doesn't", "doesnt")
        .replace("don't", "dont")
        .replace("won't", "wont")
        .replace("\u2018", "'")
        .replace("\u2019", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("\u2013", "-")
        .replace("\u2014", "-")
        .replace("\u2192", "->")
        .replace("\u00d7", "x")
        .replace("\u211d", "R")
        .replace("\u00e0", "->")
    )
    text = text.replace(" the have ", " they have ").replace(" nd returns ", " and returns ")
    text = re.sub(r"[^a-z0-9]+", " ", text.lower()).strip()
    text = text.replace("can t", "cant").replace("doesn t", "doesnt")
    text = text.replace("don t", "dont").replace("won t", "wont")
    text = re.sub(r"\badvantages\b", "advantage", text)
    text = re.sub(r"\bdisadvantages\b", "disadvantage", text)
    return text


PROOF_STOPWORDS = {
    "a", "an", "and", "are", "as", "be", "by", "can", "e", "eg", "for", "from",
    "i", "ie", "in", "is", "it", "of", "on", "or", "so", "the", "to", "with",
}

NOTE_FILE_ALIASES = {
    ("Computer organisation", "The meaning of the stored program concept"): "The stored program concept",
    ("Computer organisation", "The Fetch-Execute cycle and the role of registers within it"): "The FE cycle",
}


def proof_tokens(text: str) -> set[str]:
    return {tok for tok in normalise(text).split() if tok not in PROOF_STOPWORDS}


def text_covers(source: str, candidate_norm: str, candidate_tokens: set[str]) -> bool:
    source_norm = normalise(source)
    if not source_norm:
        return True
    if source_norm in candidate_norm:
        return True
    source_tokens = proof_tokens(source)
    return bool(source_tokens) and source_tokens <= candidate_tokens


def note_detail_cells(row) -> list[str]:
    cells = [cell.text.strip() for cell in row.cells]
    if not any(cells):
        return []
    if len(cells) >= 3 and cells[2].strip():
        return [cells[2]]
    if len(cells) >= 2:
        return [cells[1]]
    return cells


def table_cells(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [normalise(cell) for cell in s.split("|")]


def table_separator(cells: list[str]) -> bool:
    return bool(cells) and all(c and set(c) <= set("-: ") for c in cells)


def note_table_rows(text: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in text.splitlines():
        if not line.lstrip().startswith("|"):
            continue
        cells = table_cells(line)
        if cells and not table_separator(cells):
            rows.append(cells)
    return rows


def note_table_blocks(text: str) -> list[list[list[str]]]:
    blocks: list[list[list[str]]] = []
    cur: list[list[str]] = []
    for line in text.splitlines():
        if not line.lstrip().startswith("|"):
            if cur:
                blocks.append(cur)
                cur = []
            continue
        cells = table_cells(line)
        if cells and not table_separator(cells):
            cur.append(cells)
    if cur:
        blocks.append(cur)
    return blocks


def markdown_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        cells = table_cells(line)
        if cells and not table_separator(cells):
            rows.append(cells)
    return rows


def cell_covers(source: str, candidate: str) -> bool:
    source = normalise(source)
    candidate = normalise(candidate)
    if not source:
        return True
    if source == candidate or source in candidate:
        return True
    source_tokens = proof_tokens(source)
    candidate_tokens = proof_tokens(candidate)
    return bool(source_tokens) and source_tokens <= candidate_tokens


def row_present(source_cells: list[str], candidate_rows: list[list[str]]) -> bool:
    for row in candidate_rows:
        if len(row) < len(source_cells):
            continue
        if all(cell_covers(cell, row[i]) for i, cell in enumerate(source_cells) if normalise(cell)):
            return True
    for start in range(len(candidate_rows)):
        for end in range(start + 2, len(candidate_rows) + 1):
            if rows_cover(source_cells, candidate_rows[start:end]):
                return True
    return False


def rows_cover(source_cells: list[str], candidate_rows: list[list[str]]) -> bool:
    if not candidate_rows:
        return False
    cols = max(len(source_cells), max(len(row) for row in candidate_rows))
    combined = []
    for col in range(cols):
        combined.append(" ".join(row[col] for row in candidate_rows if col < len(row)))
    return all(
        cell_covers(cell, combined[i])
        for i, cell in enumerate(source_cells)
        if i < len(combined) and normalise(cell)
    )


def table_block_present(source_rows: list[list[str]], candidate_blocks: list[list[list[str]]]) -> bool:
    if not source_rows:
        return True
    for block in candidate_blocks:
        row = 0
        matched = 0
        while matched < len(source_rows) and row < len(block):
            found = False
            for end in range(row + 1, len(block) + 1):
                if rows_cover(source_rows[matched], block[row:end]):
                    row = end
                    matched += 1
                    found = True
                    break
            if not found:
                break
        if matched == len(source_rows):
            return True
    return False


def nested_tables(cell) -> list[Table]:
    return [block for block in iter_block_items(cell) if isinstance(block, Table)]


def main() -> int:
    paths = [DOCX_DIR / name for name in DOCX_NAMES]
    missing_docs = [str(path) for path in paths if not path.exists()]
    if missing_docs:
        raise AssertionError("source DOCX missing:\n" + "\n".join(missing_docs))
    notes_text = "\n".join(path.read_text(errors="ignore") for path in NOTES.rglob("*.txt"))
    notes_norm = normalise(notes_text)
    notes_tokens = proof_tokens(notes_text)
    ignored = {
        "below",
        "extra info for you codex to add to the text files avoiding duplication",
        "notes details of vocab and knowledge",
        "syllabus content",
        "syll ref",
    }
    missing: list[str] = []
    for path in paths:
        doc = Document(path)
        expected_names = DOC_TABLE_NAMES.get(path.stem, [])
        if expected_names and len(doc.tables) != len(expected_names):
            missing.append(f"{path.name}: expected {len(expected_names)} source tables, found {len(doc.tables)}")
        for block in iter_block_items(doc):
            if not isinstance(block, Paragraph):
                continue
            for raw in clean_text(block.text).splitlines():
                line = raw.strip()
                norm = normalise(line)
                if len(norm) < 3 or norm in ignored:
                    continue
                if not text_covers(line, notes_norm, notes_tokens):
                    missing.append(f"{path.name}: top-level paragraph missing: {line}")
        for table in doc.tables:
            for row in table.rows:
                for cell in note_detail_cells(row):
                    for raw in cell.splitlines():
                        line = raw.strip()
                        norm = normalise(line)
                        if len(norm) < 3 or norm in ignored:
                            continue
                        if not text_covers(line, notes_norm, notes_tokens):
                            missing.append(f"{path.name}: {line}")
        for table, table_name in zip(doc.tables, expected_names):
            note_name = NOTE_FILE_ALIASES.get((path.stem, table_name), table_name)
            note_path = NOTES / path.stem / f"{slug_safe(note_name)}.txt"
            if not note_path.exists():
                missing.append(f"{path.name}:{table_name}: expected note file missing")
                continue
            if not note_path.read_text(errors="ignore").strip():
                missing.append(f"{path.name}:{table_name}: expected note file is empty")
                continue
            table_note_norm = normalise(note_path.read_text(errors="ignore"))
            table_note_tokens = proof_tokens(note_path.read_text(errors="ignore"))
            detail_col = table_detail_column(table)
            for row in table.rows:
                for cell in note_detail_cells(row):
                    for raw in cell.splitlines():
                        line = raw.strip()
                        norm = normalise(line)
                        if len(norm) < 3 or norm in ignored:
                            continue
                        if not text_covers(line, table_note_norm, table_note_tokens):
                            missing.append(f"{path.name}:{table_name}: misplaced/missing from table note: {line}")
            rows_with_nested = []
            for row_index, row in enumerate(table.rows[1:], 2):
                if detail_col >= len(row.cells):
                    continue
                blocks = nested_tables(row.cells[detail_col])
                if blocks:
                    rows_with_nested.append((row_index, blocks))
            if not rows_with_nested:
                continue
            note_text = note_path.read_text(errors="ignore")
            table_rows = note_table_rows(note_text)
            table_blocks = note_table_blocks(note_text)
            for row_index, nested_blocks in rows_with_nested:
                for nested_index, nested in enumerate(nested_blocks, 1):
                    source_rows = markdown_rows(markdown_table(nested))
                    if not table_block_present(source_rows, table_blocks):
                        missing.append(
                            f"{path.name}:{table_name}: nested table block missing "
                            f"(source row {row_index}, nested table {nested_index})"
                        )
                    for raw in markdown_table(nested):
                        cells = table_cells(raw)
                        if not cells or table_separator(cells):
                            continue
                        if not row_present(cells, table_rows):
                            missing.append(
                                f"{path.name}:{table_name}: nested table row missing "
                                f"(source row {row_index}, nested table {nested_index}): {raw}"
                            )
    if missing:
        raise AssertionError("\n".join(missing[:80]))
    print(f"OK notes docx coverage: {len(paths)} source docs")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
